"""One-off generator for the HALO Clean Architecture / Riverpod progress report."""

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor, Emu, Inches

OUT_DIR = Path(__file__).resolve().parent
DOCX_PATH = OUT_DIR / "HALO_Clean_Architecture_Riverpod_Progress_Report.docx"

NAVY = RGBColor(0x1B, 0x2A, 0x4A)
ACCENT = RGBColor(0x5B, 0x3F, 0xA3)
MUTED = RGBColor(0x4A, 0x55, 0x68)
GREEN = RGBColor(0x1B, 0x7A, 0x4A)
AMBER = RGBColor(0xB5, 0x6E, 0x00)
RED = RGBColor(0xB4, 0x23, 0x18)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_BG = "1B2A4A"
ROW_ALT = "F4F1FB"
DONE_BG = "E6F4EA"
PARTIAL_BG = "FFF4D6"
TODO_BG = "FCE8E6"


def set_run_font(run, name="Calibri", size=11, bold=False, color=None, italic=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_cell_borders(cell, color="D0D5DD"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def set_cell_text(cell, text, *, bold=False, size=10, color=NAVY, align="left", fill=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.08
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    if fill:
        shade_cell(cell, fill)
    set_cell_borders(cell)
    cell.vertical_alignment = 1  # center


def add_page_number(paragraph):
    run = paragraph.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instr)
    run._r.append(fldChar2)
    set_run_font(run, size=9, color=MUTED)


def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = NAVY if level == 1 else ACCENT
        run.font.name = "Calibri"
    h.paragraph_format.space_before = Pt(16 if level == 1 else 12)
    h.paragraph_format.space_after = Pt(6)
    return h


def add_body(doc, text, *, bold=False, italic=False, size=11, space_after=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic, color=NAVY)
    return p


def add_bullet(doc, text, *, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.12
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, size=11, bold=True, color=NAVY)
        r2 = p.add_run(text)
        set_run_font(r2, size=11, color=NAVY)
    else:
        r = p.add_run(text)
        set_run_font(r, size=11, color=NAVY)
    return p


def add_callout(doc, title, body, fill="F4F1FB"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    set_cell_borders(cell, "C4B5E0")
    cell.text = ""
    p1 = cell.paragraphs[0]
    p1.paragraph_format.space_before = Pt(8)
    p1.paragraph_format.space_after = Pt(2)
    r = p1.add_run(title)
    set_run_font(r, size=11, bold=True, color=ACCENT)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(8)
    p2.paragraph_format.line_spacing = 1.15
    r2 = p2.add_run(body)
    set_run_font(r2, size=10.5, color=NAVY)
    doc.add_paragraph()


def make_table(doc, headers, rows, col_widths=None, status_col=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True, size=10, color=WHITE, align="center", fill=HEADER_BG)
    for r_i, row in enumerate(rows):
        alt = ROW_ALT if r_i % 2 else "FFFFFF"
        for c_i, val in enumerate(row):
            fill = alt
            color = NAVY
            bold = False
            if status_col is not None and c_i == status_col:
                low = str(val).lower()
                if "done" in low or "complete" in low or low.startswith("yes"):
                    fill, color, bold = DONE_BG, GREEN, True
                elif "partial" in low or "in progress" in low:
                    fill, color, bold = PARTIAL_BG, AMBER, True
                elif "not" in low or "none" in low or low.startswith("no"):
                    fill, color, bold = TODO_BG, RED, True
            align = "center" if c_i != 0 else "left"
            set_cell_text(table.cell(r_i + 1, c_i), str(val), size=9.5, color=color, bold=bold, align=align, fill=fill)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)

    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hr = hp.add_run("HALO  ·  Architecture Progress Report  ·  Confidential")
    set_run_font(hr, size=9, color=MUTED, italic=True)

    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run("Prepared 1 September 2026  ·  Page ")
    set_run_font(fr, size=9, color=MUTED)
    add_page_number(fp)

    # ----- Title -----
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.LEFT
    t.paragraph_format.space_before = Pt(12)
    r = t.add_run("HALO")
    set_run_font(r, size=14, bold=True, color=ACCENT)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    tr = title.add_run("Clean Architecture & Riverpod\nProgress Report")
    set_run_font(tr, size=26, bold=True, color=NAVY)

    sub = doc.add_paragraph()
    sr = sub.add_run("Status of the architecture migration on branch sahyog versus current main, as of 1 September 2026.")
    set_run_font(sr, size=12, italic=True, color=MUTED)

    meta = [
        ("Product", "HALO (Flutter + Firebase)"),
        ("Repository", "classic_1_backup"),
        ("Architecture branch", "sahyog  (this is where the work lives)"),
        ("Current checkout", "main  (no Riverpod, no features/auth)"),
        ("Report date", "1 September 2026"),
        ("Scope", "Honest inventory — not a launch plan"),
    ]
    make_table(doc, ["Item", "Value"], meta, col_widths=[5.5, 11.5])

    add_callout(
        doc,
        "Read this first",
        "Clean Architecture + Riverpod is not on main. Checking out main (or building from main) shows the old login in lib/main.dart and lib/Bottom Pages/login_page.dart, with Firebase called from widgets. All numbers below for “done” work refer to branch sahyog unless stated otherwise. sahyog is also not merged: main is 1 commit ahead and 29 commits behind origin/main in the working tree snapshot used for this report.",
        fill="FFF4D6",
    )

    # ----- 1 Executive summary -----
    add_heading_styled(doc, "1. Executive summary", 1)
    add_body(
        doc,
        "The app is not “already on Clean Architecture.” One vertical slice is: authentication and HALO onboarding. That slice has a real domain contract, a Firebase data layer, Riverpod controllers, and a session-based gate. Everything else — home, explore, reels, chat, search, stories, notifications, upload, and the three giant profile pages — still talks to Firebase from UI or singleton services.",
    )

    make_table(
        doc,
        ["Score", "Value", "How it is measured"],
        [
            ["Overall app migrated", "~10–15%", "Dart files using Riverpod (16 of 167) and lines in architecture folders vs ~44,400 total"],
            ["Auth + onboarding slice", "~85–90%", "Domain, data, presentation, gate, tests for session rules"],
            ["Riverpod in the UI", "11 of ~87 screens", "ConsumerWidget / ConsumerStatefulWidget vs other widget files"],
            ["Clean feature modules", "1 of ~10", "Only lib/features/auth exists. No features/feed, chat, reels, profile-write"],
            ["main branch progress", "0%", "flutter_riverpod is not a dependency; no lib/features folder"],
            ["Automated tests", "Thin", "3 test files; session mapper covered; no Home/Explore/Reels tests"],
        ],
        col_widths=[5.0, 3.5, 8.5],
    )

    add_body(
        doc,
        "Bottom line: do not rewrite the remaining 85% before a beta. Keep the strangler pattern — finish merging sahyog auth onto a stable branch, then migrate Home / feed / upload the same way, one slice at a time.",
        bold=False,
    )

    # ----- 2 What “done” means -----
    add_heading_styled(doc, "2. What “Clean Architecture + Riverpod” means in this project", 1)
    add_body(
        doc,
        "The intended pattern (already applied to auth) is the standard Flutter feature-first layout, not a textbook enterprise stack with use-case classes for every button.",
    )
    add_bullet(doc, "pure Dart types and rules. No Firebase imports.", bold_prefix="Domain — ")
    add_bullet(doc, "implements the domain contract. Firebase Auth, Firestore, Google, Apple, and phone OTP live only here.", bold_prefix="Data — ")
    add_bullet(doc, "pages and widgets. They call Riverpod notifiers. They do not call FirebaseAuth.instance or FirebaseFirestore.instance.", bold_prefix="Presentation — ")
    add_bullet(doc, "Provider / StreamProvider / StateNotifierProvider. ProviderScope wraps the app in main.dart.", bold_prefix="Riverpod — ")
    add_bullet(doc, "one Session stream decides Login vs Verify Email vs Category vs Profile form vs Home.", bold_prefix="Gate — ")

    add_body(doc, "A feature is counted as complete only if all four are true: domain contract, data implementation, Riverpod wiring, and UI free of Firebase.")

    # ----- 3 Completed -----
    add_heading_styled(doc, "3. What is complete (branch sahyog)", 1)

    add_heading_styled(doc, "3.1 Authentication feature module", 2)
    add_body(
        doc,
        "lib/features/auth is a real three-layer module (~3,150 lines, 16 Dart files). The UI for login, email signup, phone OTP, verify email, and forgot password goes through repositories and Riverpod. New email users get a Firebase account, a verification email, then category + profile. Google / Phone / Apple authenticate first, then the same onboarding if the HALO profile is incomplete. Existing users with onboardingCompleted (or a legacy username + accountType) go to Home.",
    )

    make_table(
        doc,
        ["Layer", "Files", "Role"],
        [
            ["Domain", "auth_repository.dart, session_mapper.dart, phone_otp_session.dart", "Contract: sign in/up, OTP, Google, Apple, phone, verification, password reset, username check, complete onboarding, clear account type"],
            ["Core session", "lib/core/session.dart", "SessionStatus: loading, loggedOut, emailVerificationRequired, onboardingRequired, authenticated"],
            ["Data", "firebase_auth_repository.dart (~553 lines)", "Only place auth talks to Firebase. Writes users/{uid}. Does not store passwords. Strips secret keys on profile merge"],
            ["Riverpod", "session_controller.dart, phone_auth_controller.dart", "authRepositoryProvider, sessionProvider, authActionProvider, onboardingControllerProvider, phoneAuthControllerProvider"],
            ["Gate", "auth_gate.dart, onboarding_gate.dart", "Single router from session. Category page if type is null; matching profile form if type is set"],
            ["Pages", "login, email_signup, phone_login, verify_email", "Consumer widgets. No Firebase in these pages"],
            ["Onboarding UI", "onboarding_ui.dart + Category + 3 create-account forms", "Forms submit via onboardingControllerProvider.completeOnboarding"],
        ],
        col_widths=[3.5, 6.5, 7.0],
    )

    add_heading_styled(doc, "3.2 Auth behaviours that are implemented", 2)
    add_bullet(doc, "Email/password, username, or mobile identifier login via the repository.")
    add_bullet(doc, "Email signup creates the account, sends verification, stubs users/{uid} with onboardingCompleted: false.")
    add_bullet(doc, "Verify Email page: reload, resend, sign out. Password users cannot skip verification.")
    add_bullet(doc, "Forgot password via sendPasswordResetEmail on the repository.")
    add_bullet(doc, "Google and Apple sign-in in the data layer (Apple is implemented in code; iOS project config is still incomplete).")
    add_bullet(doc, "Phone OTP with verification-id refresh; UI pops when session leaves loggedOut.")
    add_bullet(doc, "Category pick writes accountType only. Back / “Change type” calls clearAccountType so the user is not stuck on Wellness/Guru.")
    add_bullet(doc, "Profile onboarding (Aspirant / Guru / Wellness) no longer creates a second Firebase account and no longer writes passwords.")
    add_bullet(doc, "Username uniqueness checked in the data layer.")
    add_bullet(doc, "Centered toasts (halo_toast) and white/modern onboarding chrome.")

    add_heading_styled(doc, "3.3 Riverpod providers that exist", 2)
    make_table(
        doc,
        ["Provider", "Type", "Used by"],
        [
            ["authRepositoryProvider", "Provider<AuthRepository>", "All auth controllers"],
            ["sessionProvider", "StreamProvider<Session>", "OnboardingGate (app router)"],
            ["authActionProvider", "StateNotifierProvider", "Login, signup, verify, Google, Apple, sign-out, password reset"],
            ["onboardingControllerProvider", "StateNotifierProvider", "Category + three profile forms"],
            ["phoneAuthControllerProvider", "StateNotifierProvider", "Phone OTP page"],
            ["profileRepositoryProvider", "Provider<ProfileRepository>", "Profile view (partial)"],
            ["profileStreamProvider", "StreamProvider", "Profile view (partial)"],
            ["unifiedProfileDataProvider", "StreamProvider.family", "Dynamic profile page (partial)"],
        ],
        col_widths=[6.2, 5.0, 5.8],
    )

    add_heading_styled(doc, "3.4 Tests that exist for this slice", 2)
    add_body(
        doc,
        "test/profile_type_test.dart covers session mapping: unverified email is gated; missing HALO doc goes to category selection; onboardingCompleted or legacy username+accountType goes to authenticated; Google/phone users are not forced through email verification. test/profile_widgets_test.dart and test/widget_test.dart are small smoke tests. There are no integration tests for login, OTP, or Home.",
    )

    # ----- 4 Partial -----
    add_heading_styled(doc, "4. What is only partial", 1)

    add_heading_styled(doc, "4.1 Profile viewing (lib/screens/profile)", 2)
    add_body(
        doc,
        "A modular profile UI was started: 44 files, ~2,380 lines, with ProfileKind, a thin repository barrel, and a Riverpod stream for ProfileData. That is folder structure plus some extraction — not a finished Clean Architecture feature. Several profile widgets still query Firestore directly (recent posts, online dot). The old Profile Pages were not deleted.",
    )
    make_table(
        doc,
        ["Old profile page (still in repo)", "Approx. lines", "Riverpod / domain"],
        [
            ["lib/Profile Pages/aspirant_profile_page.dart", "3,351", "No"],
            ["lib/Profile Pages/guru_profile_page.dart", "5,522", "No"],
            ["lib/Profile Pages/wellness_profile_page.dart", "4,188", "No"],
            ["lib/Profile Pages/edit_profile_sections.dart", "Firestore × 18 call sites", "No"],
        ],
        col_widths=[8.5, 4.0, 4.5],
        status_col=2,
    )
    add_body(doc, "Those three pages alone are ~13,000 lines of AI-generated UI. They are the largest remaining refactor, and they should not be rewritten before Aspirant beta testers are using the app.")

    add_heading_styled(doc, "4.2 Onboarding profile forms", 2)
    add_body(
        doc,
        "createaspirantaccount.dart (~782 lines), createguruaccount.dart (~953), createwellnessaccount.dart (~830) now submit through Riverpod and no longer call createUserWithEmailAndPassword. They are still oversized form widgets, not domain models. Guru/Wellness certification fields can still store local file paths. That is a data-quality bug, not an architecture win.",
    )

    add_heading_styled(doc, "4.3 Services folder", 2)
    add_body(
        doc,
        "lib/services has 34 Dart files / ~3,300 lines (feed, explore, reels, stories, upload, follow, save, search, video pipeline). These are singleton helpers, not domain repositories, and the screens that call them are not Riverpod consumers. This is a useful extraction for ranking and video memory — it is not Clean Architecture.",
    )

    # ----- 5 Not started -----
    add_heading_styled(doc, "5. What has not been migrated", 1)
    add_body(doc, "On sahyog, about 54 Dart files still reference FirebaseAuth.instance or FirebaseFirestore.instance. The high-impact screens:")
    make_table(
        doc,
        ["Area", "Main files", "Firebase in UI", "Riverpod", "Clean feature folder"],
        [
            ["Home / feed", "HomePage.dart (~1,743 lines)", "17 call sites", "No", "No"],
            ["Explore", "ExplorePage.dart (~1,522 lines)", "9", "No", "No"],
            ["Reels", "reels_feed.dart (~524 lines)", "7", "No", "No"],
            ["Chat", "chat_screen / chat_list / user_list", "12+", "No", "No"],
            ["Search", "SearchPage.dart", "7", "No", "No"],
            ["Notifications", "NotificationPage.dart", "4", "No", "No"],
            ["Stories", "story_viewer, user_story_pager", "7", "No", "No"],
            ["Upload / posts", "AddPostPage, upload_service", "Yes", "No", "No"],
            ["Settings / privacy", "SettingsPage, PrivacySettingsPage", "4", "No", "No"],
            ["Interest selection", "interest_selection_page.dart", "2", "No", "No"],
            ["Duplicate home/login", "home_page.dart, leftover startup_router", "Yes", "Leftover", "No"],
        ],
        col_widths=[3.4, 5.2, 2.6, 2.2, 3.6],
        status_col=3,
    )

    add_heading_styled(doc, "5.1 main branch (current checkout)", 2)
    add_body(
        doc,
        "The workspace on 1 September 2026 is on main. That tree has 74 Dart files, no flutter_riverpod dependency, no lib/features, and login still mixed into lib/main.dart. Building an APK from main does not include the architecture work described in sections 3–4. Merge or cherry-pick sahyog before treating auth as “done” on the branch you ship.",
    )

    # ----- 6 Scorecard -----
    add_heading_styled(doc, "6. Feature scorecard", 1)
    make_table(
        doc,
        ["Feature", "Domain", "Data repo", "Riverpod", "UI free of Firebase", "Tests", "Verdict"],
        [
            ["Email / Google / Phone / Apple login", "Yes", "Yes", "Yes", "Yes (login pages)", "Session mapper", "Done"],
            ["Email verification gate", "Yes", "Yes", "Yes", "Yes", "Yes", "Done"],
            ["Password reset", "Yes", "Yes", "Yes", "Mostly", "No", "Done"],
            ["Session routing / AuthGate", "Yes", "Yes", "Yes", "Yes", "Yes", "Done"],
            ["Category pick + back/clear type", "Yes", "Yes", "Yes", "Yes", "Partial", "Done"],
            ["Aspirant / Guru / Wellness onboarding save", "Partial", "Yes", "Yes", "Mostly", "No", "Partial"],
            ["Profile view (new screens/profile)", "Thin", "Thin", "Partial", "No", "Smoke", "Partial"],
            ["Legacy profile pages", "No", "No", "No", "No", "No", "Not started"],
            ["Home / feed", "No", "Service only", "No", "No", "No", "Not started"],
            ["Explore / search ranking", "No", "Service only", "No", "No", "No", "Not started"],
            ["Reels / video pipeline", "No", "Services", "No", "No", "No", "Not started"],
            ["Stories", "No", "Service only", "No", "No", "No", "Not started"],
            ["Chat", "No", "chat_service", "No", "No", "No", "Not started"],
            ["Upload / Cloudinary", "No", "Services", "No", "No", "No", "Not started"],
            ["Notifications", "No", "No", "No", "No", "No", "Not started"],
            ["iOS Sign in with Apple + plist", "Code exists", "Code exists", "Yes", "N/A", "No", "Partial"],
        ],
        col_widths=[4.0, 1.8, 2.2, 1.8, 2.4, 2.0, 2.2],
        status_col=6,
    )

    # ----- 7 Inventory -----
    add_heading_styled(doc, "7. Size inventory (sahyog)", 1)
    make_table(
        doc,
        ["Bucket", "Files", "Approx. lines", "Share of lib/"],
        [
            ["All Dart under lib/", "167", "44,393", "100%"],
            ["features/auth (Clean + Riverpod)", "16", "3,148", "7%"],
            ["core/ (session, theme, toast, splash)", "4", "379", "<1%"],
            ["screens/profile (modular UI, mixed)", "44", "2,380", "5%"],
            ["services/ (singletons, not Riverpod UI)", "34", "3,308", "7%"],
            ["Legacy Profile Pages", "4", "~13,060", "~29%"],
            ["Home + Explore", "2", "~3,265", "~7%"],
            ["Riverpod imports anywhere", "16 files", "—", "10% of files"],
            ["Consumer widgets", "11 files", "—", "~13% of UI files"],
        ],
        col_widths=[7.5, 2.5, 3.5, 3.5],
    )

    add_body(
        doc,
        "Architecture folders look larger than the real migration. screens/profile and services/ inflate the “structured code” count. The only slice that meets the definition in section 2 is auth.",
        italic=True,
    )

    # ----- 8 Leftovers / risks -----
    add_heading_styled(doc, "8. Known leftovers and risks", 1)
    add_bullet(doc, "is leftover and still touches Firebase. Interest selection is no longer an auth gate.", bold_prefix="startup_router.dart ")
    add_bullet(doc, "Duplicate HomePage.dart vs home_page.dart and old login paths can confuse merges.")
    add_bullet(doc, "Guru/Wellness certification images may store local paths instead of uploaded URLs.")
    add_bullet(doc, "Email verification depends on Firebase (noreply@halo-fb212.firebaseapp.com). Spam and rate limits are product issues, not architecture gaps.")
    add_bullet(doc, "iOS: GoogleService-Info.plist / Apple Sign-In capability / display name were incomplete on main. Repository methods exist on sahyog; store listing and signing are not done.")
    add_bullet(doc, "sahyog unmerged. Shipping from main silently drops the entire auth rewrite.")
    add_bullet(doc, "No crash/jank budget, no CI test gate beyond the three local tests.")

    # ----- 9 Recommended next slices -----
    add_heading_styled(doc, "9. Recommended next slices (do not rewrite the whole app)", 1)
    add_body(doc, "Use the same pattern as auth. One feature folder, one repository, a few providers, then delete Firebase from that screen.")
    make_table(
        doc,
        ["Order", "Slice", "Why this next", "Difficulty"],
        [
            ["0", "Merge sahyog auth onto the branch you ship; smoke-test login on a device", "Otherwise the work in this report is not in the APK", "Medium (merge conflict risk)"],
            ["1", "Sign-out / session on Home and Settings (stop FirebaseAuth.instance there)", "Auth is already the source of truth; Home still bypasses it", "Low"],
            ["2", "Feed / Home read path → features/feed", "Highest user time, 17 Firebase call sites", "High"],
            ["3", "Upload / new post → features/upload", "Crashes and duplicate posts hurt beta more than profile chrome", "High"],
            ["4", "Aspirant profile view + edit only", "Needed for Aspirant-only beta; leave Guru/Wellness giants", "Medium"],
            ["5", "Chat list + thread", "Direct Firebase in several chat files", "Medium"],
            ["6", "Explore / reels / stories", "Already have ranking services; wrap them, do not rewrite ranking", "High"],
            ["7", "Guru then Wellness profile pages", "Only after Aspirant beta is stable", "High"],
            ["8", "iOS TestFlight (plist, Apple Sign-In, signing)", "Not a “minor toggle”", "Medium–High"],
        ],
        col_widths=[1.8, 5.5, 7.2, 2.5],
    )

    add_heading_styled(doc, "9.1 Definition of done for the next slice", 2)
    add_bullet(doc, "lib/features/<name>/{domain,data,presentation} exists.")
    add_bullet(doc, "Widgets in that slice do not import firebase_auth or cloud_firestore.")
    add_bullet(doc, "At least one unit test for the mapper or repository contract.")
    add_bullet(doc, "Old duplicate file deleted or reduced to a re-export.")

    # ----- 10 Comparison -----
    add_heading_styled(doc, "10. main vs sahyog at a glance", 1)
    make_table(
        doc,
        ["Topic", "main (current checkout)", "sahyog (architecture work)"],
        [
            ["Dart files in lib/", "74", "167"],
            ["flutter_riverpod", "Not in pubspec", "^2.6.1"],
            ["ProviderScope", "No", "Yes, in main.dart"],
            ["lib/features/auth", "Missing", "16 files, 3 layers"],
            ["Session gate", "Ad-hoc in main / login_page", "OnboardingGate + SessionStatus"],
            ["Login Firebase in widgets", "Yes", "No (login pages)"],
            ["Onboarding creates 2nd account", "Yes (legacy forms)", "No — completeProfileOnboarding on same uid"],
            ["Profile module", "Monolith pages only", "New screens/profile + old pages both present"],
            ["Tests", "widget_test smoke", "session mapper + profile smoke"],
        ],
        col_widths=[4.2, 6.4, 6.4],
    )

    # ----- 11 Conclusion -----
    add_heading_styled(doc, "11. Conclusion", 1)
    add_body(
        doc,
        "Clean Architecture and Riverpod are started, not finished. Auth is the proof that the pattern works: domain contract, Firebase isolated in data, Riverpod controllers, and a single session gate. That is roughly one-tenth of the app by files, and the most important tenth for a correct login.",
    )
    add_body(
        doc,
        "Calling the project “architected” would be inaccurate while Home, Explore, Reels, Chat, and ~13,000 lines of profile UI still call Firebase. The fastest path to a beta is to protect the auth slice, merge it, and repeat the same pattern on feed and upload — not to pause launch for a full rewrite.",
    )

    add_heading_styled(doc, "Appendix A — Auth file list (sahyog)", 2)
    auth_files = [
        "lib/core/session.dart",
        "lib/core/halo_toast.dart",
        "lib/core/halo_theme.dart",
        "lib/core/halo_splash.dart",
        "lib/features/auth/domain/auth_repository.dart",
        "lib/features/auth/domain/session_mapper.dart",
        "lib/features/auth/domain/phone_otp_session.dart",
        "lib/features/auth/data/firebase_auth_repository.dart",
        "lib/features/auth/presentation/session_controller.dart",
        "lib/features/auth/presentation/phone_auth_controller.dart",
        "lib/features/auth/presentation/auth_gate.dart",
        "lib/features/auth/presentation/onboarding_gate.dart",
        "lib/features/auth/presentation/onboarding_ui.dart",
        "lib/features/auth/presentation/pages/login_page.dart",
        "lib/features/auth/presentation/pages/email_signup_page.dart",
        "lib/features/auth/presentation/pages/phone_login_page.dart",
        "lib/features/auth/presentation/pages/verify_email_page.dart",
        "lib/features/auth/presentation/widgets/login_button.dart",
        "lib/features/auth/presentation/widgets/legal_dialogs.dart",
        "lib/Category/categorypage.dart",
        "lib/Category/createaspirantaccount.dart",
        "lib/Category/createguruaccount.dart",
        "lib/Category/createwellnessaccount.dart",
        "test/profile_type_test.dart",
    ]
    for f in auth_files:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(f)
        set_run_font(r, name="Consolas", size=9.5, color=NAVY)

    add_heading_styled(doc, "Appendix B — Method notes", 2)
    add_body(
        doc,
        "Counts were taken from git show against branch sahyog on 1 September 2026 (PowerShell Measure-Object -Line on each Dart file). “Call sites” are git grep counts of FirebaseAuth.instance and FirebaseFirestore.instance. Percentages are rounded. This document does not include a time/cost estimate; those live in the separate beta roadmap discussion.",
        size=10,
        italic=True,
    )

    doc.save(DOCX_PATH)
    print(f"Wrote {DOCX_PATH}")


if __name__ == "__main__":
    build()
